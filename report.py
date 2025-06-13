import pandas as pd
import warnings
from datetime import *
import copy
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, RGBColor  
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def get_dates_in_week(year, week):
    #It returns all the working days in a specific week
    start_date = datetime.fromisocalendar(year, week, 1).date()
    return [start_date + timedelta(days=i) for i in range(5)]


class Cosmethics():

    def __init__(self,table):
        self.table=table
    def set_cell_bg_color(self,cell, color_hex: str):
        """
        Set the color of cell background.
        color_hex: string ex. 'D9E1F2' (without #)
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

    def decorate_table(self, *,
                    header_rows=2,
                    header_bg='D9E1F2',
                    font_name='Arial',
                    font_size=10,
                    bold_headers=True,
                    align_center=True):
        """
        It applies a decoration to a Word table.
        Args:
            table: Table object of python-docx
            header_rows: number of header rows
            header_bg: Header background color (HEX without #)
            font_name: Font name
            font_size: Font size
            bold_headers: If True, it returns bold header
            align_center: If True, it centers the text
        """
        for row_idx, row in enumerate(self.table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    if bold_headers and row_idx < header_rows:
                        run.bold = True
                    if align_center:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if row_idx < header_rows:
                    self.set_cell_bg_color(cell, header_bg)



class Report:
    def __init__(self,document_path,year,week):
        self.week=week
        self.doc = Document()
        section = self.doc.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(10.3)   # Default A4: 8.27
        section.page_height = Inches(12)  # Default A4: 11.7

        # Header
        header = section.header
        header_paragraph = header.paragraphs[0]
        run_header = header_paragraph.add_run("LCAJT Weekly Report")  
        run_header.font.name = 'Calibri'
        run_header.font.size = Pt(11)
        run_header.font.color.rgb = RGBColor(100, 100, 100)  
        

        #First Title
        #title =self.doc.add_heading(level=1)
        title = self.doc.add_paragraph()
        self.run=title.add_run(f"GBTS Weekly Report CW {str(self.week)}")
        self.run.font.name ='Calibri'
        self.run.font.size =Pt(28)
        self.run.font.color.rgb = RGBColor(0, 0, 0)

        #Subtitle
        days=get_dates_in_week(year,week)
        subtitle = self.doc.add_paragraph()
        #run_subtitle = subtitle.add_run(f"{str(days[0])} - {str(days[4])}")  
        run_subtitle = subtitle.add_run(f"{days[0].strftime("%d %b %Y")} - {days[4].strftime("%d %b %Y")}")
        run_subtitle.font.name = 'Calibri'
        run_subtitle.font.size = Pt(18)
        run_subtitle.font.color.rgb = RGBColor(0, 0, 0)
        self.name=f"{document_path}Weekly Report - Week {week}.docx"

        #Footer
        
        self.add_footer_image(os.path.join("images", "ajt_official.png"))
        self.add_page_number_total()
        
        
        

    def convert_in_landscape(self):
        section=self.doc.add_section(start_type=1)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    def convert_in_portrait(self):
        section=self.doc.add_section(start_type=1)
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = section.page_height, section.page_width

    def new_paragraph(self,title,level=1,font_size=16,r_color=54,g_color=95,b_color=145):
        #paragraph = self.doc.add_paragraph()
        paragraph=self.doc.add_heading(level=level)
        run = paragraph.add_run(title)
        run.font.name = 'Cambria'
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(r_color,g_color,b_color)
        run.bold = True
        #paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.left_indent = Inches(0)
        paragraph.paragraph_format.first_line_indent = Inches(0)
        return paragraph

    def add_footer_image(self, image_path):
        """
        Inserisce un'immagine nel piè di pagina, allineata a sinistra.
        """
        section = self.doc.sections[0]
        footer = section.footer

        # Se esiste già un paragrafo nel footer, usalo, altrimenti creane uno
        if footer.paragraphs:
            paragraph_logo = footer.paragraphs[0]
        else:
            paragraph_logo = footer.add_paragraph()

        paragraph_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_logo = paragraph_logo.add_run()
        run_logo.add_picture(image_path, width=Inches(2.5))

    def add_page_number_total(self):
        section = self.doc.sections[0]
        footer = section.footer
        #paragraph = footer.paragraphs[0]
        paragraph = footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  

        # "Page "
        run = paragraph.add_run()
        run.font.size = Pt(7)

        # Current page number: { PAGE }
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldText = OxmlElement('w:t')
        fldText.text = "1"

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldText)
        run._r.append(fldChar3)

        run = paragraph.add_run(" / ")
        run.font.size = Pt(7)

        # Total number of pages: { NUMPAGES }
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'NUMPAGES'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldText = OxmlElement('w:t')
        fldText.text = "1"

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldText)
        run._r.append(fldChar3)


    


    def generate_text(self,text,font_size=11,font_name='Calibri'):
        paragraph=self.doc.add_paragraph()
        description=paragraph.add_run(text)
        description.font.name=font_name
        description.font.size=Pt(font_size)

    def generate_sim_util_table(self,data:dict):
        table=self.doc.add_table (rows=10,cols=11)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        devices = ["FMS1", "FMS2", "PTT1", "PTT2", "PTT3", "ULTD1", "ULTD2", "LVC"]
        
        # -- headers row 0 --
        cell = table.cell(0, 0)
        cell.text = ""
        cell.vertical_alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, date in enumerate(data.keys()):
            cell = table.cell(0, 1 + i*2)
            cell.text = str(date)
            cell.merge(table.cell(0, 1 + i*2 + 1))  

        dev=table.cell(1, 0)
        dev.text = "DEVICE"  
        
        for i in range(len(data.keys())):
            table.cell(1, 1 + i*2).text = "PLANNED"
            table.cell(1, 1 + i*2 + 1).text = "COMPLETED"
        
        for row_index, device in enumerate(devices):
            table.cell(2 + row_index, 0).text = device 
            for i,date in enumerate(data.keys()):
                numbers=data[date]
                table.cell(2 + row_index, 1 + i*2).text = str(numbers[row_index][0])
                table.cell(2 + row_index, 1 + i*2+1).text = str(numbers[row_index][1])

        
        for i, date in enumerate(list(data.keys())):
            planned_col = [int(table.cell(2 + r, 1 + i*2).text) for r in range(len(devices))]
            completed_col = [int(table.cell(2 + r, 1 + i*2 + 1).text) for r in range(len(devices))]
            
            if all(v == 0 for v in planned_col + completed_col):
                # Cancella i valori della colonna
                for r in range(2, len(devices) + 2):
                    table.cell(r, 1 + i*2).text = ""
                    table.cell(r, 1 + i*2 + 1).text = ""

                # Unisci celle per messaggio
                merged_cell = table.cell(2, 1 + i*2)
                for r in range(3, len(devices) + 2):
                    #merged_cell = merged_cell.merge(table.cell(r, 1 + i*2))
                    merged_cell = merged_cell.merge(table.cell(r, 1 + i*2+1))

                merged_cell.text = ""  # svuota prima
                merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # Coloured text
                paragraph = merged_cell.paragraphs[0]
                run = paragraph.add_run("No Training")
                run.font.name = "Calibri"
                run.font.size = Pt(16)
                run.font.bold=True
                run.font.color.rgb = RGBColor(255, 255, 255)  
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.set_cell_bg_color(merged_cell,'D8D3D3')

                

        self.doc.add_paragraph()#just to have a space between the table and the next content
        return table

    def generate_generic_table(self,n_columns,data:list,header:list):
        table=self.doc.add_table (rows=len(data)+1,cols=n_columns)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i,item in enumerate(header):
            table.cell(0,i).text=str(item) 
        
        for row,items in enumerate(data):
            for col,el in enumerate(items):
                table.cell(row+1,col).text=str(el) 
        table.allow_autofit = False
        self.doc.add_paragraph()#just to have a space between the table and the next content
        return table
    
    def add_chart(self,img_chart):
        self.doc.add_picture(img_chart, width=Inches(7))

    def legend(self,title:str,rows=1,cols=1,cell_content=[]):
        table=self.doc.add_table (rows=rows,cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell=table.cell(0,0)
        cell.text=title
        cell.merge(table.cell(0,cols-1))
        for r in range(1,rows):
            for c in range(0,cols):
                table.cell(r,c).text=cell_content[r-1][c]
        return table

    

    def set_cell_bg_color(self,cell, color_hex: str):
        """
        Imposta il colore di sfondo della cella.
        color_hex: stringa es. 'D9E1F2' (senza #)
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

    def decorate_table(self, table, *, 
                   table_alignment_center=True,
                   header_rows=1,
                   header_bg='D9E1F2',
                   font_name='Arial',
                   font_size=10,
                   bold_headers=True,
                   align_center=True,
                   total_width_cm=None,
                   column_widths_cm=None,
                   alternate_row_color='EFF5FB',
                   columns_left_alignment=[] #it contains the number of columns to align at left side
                   ):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # Imposta larghezza colonna
                if column_widths_cm and col_idx < len(column_widths_cm):
                    cell.width = Cm(column_widths_cm[col_idx])
                
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # Font e allineamento
                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    if bold_headers and row_idx < header_rows:
                        run.bold = True
                    
                    if row_idx < header_rows:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # sempre centrato per intestazione
                    elif col_idx in columns_left_alignment:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif align_center:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # fallback

                # Colore intestazione
                if row_idx < header_rows:
                    self.set_cell_bg_color(cell, header_bg)
                elif header_rows < 2 and row_idx % 2 == 1:
                    self.set_cell_bg_color(cell, alternate_row_color)  # righe alternate se niente header

        if table_alignment_center:
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        else: table.alignment = WD_TABLE_ALIGNMENT.LEFT

        if total_width_cm:
            table.autofit = False
            table.allow_autofit = False
            table.preferred_width = Cm(total_width_cm)
        
        #self.shift_table_right(table,shift_cm=0.4)


    def save_documents(self):
            try:
                self.doc.save(self.name)
                print("Document saved to ",self.name)
            except Exception as saveErr:
                print(saveErr)

#print(get_dates_in_week(2025,20))

